"""
tools/cover_letter.py — Generates a personalised cover letter for each application.

Pipeline:
  1. Load base template from templates/cover_letter_base.txt
  2. Read the tailored CV content (already written by cv_customizer) for context
  3. Build a detailed Claude prompt: JD + company profile + CV content + rules
  4. Call Claude Sonnet — same model as CV, quality matters
  5. Hydrate the letter with real contact info from .env
  6. Render to a clean, professional .docx via python-docx
  7. Return the file path for the tracker

Fallback: if Claude fails, we fill the base template with known facts
(company name, job title, contact info) and write that — so the pipeline
always produces a file even in degraded mode.
"""

import os
import re
from datetime import date
from pathlib import Path
from typing import Optional

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv

from config import CANDIDATE_PROFILE, OUTPUT_CONFIG

load_dotenv()

# ---------------------------------------------------------------------------
# Paths and model
# ---------------------------------------------------------------------------
BASE_TEMPLATE_PATH = Path("templates/cover_letter_base.txt")
OUTPUT_DIR         = Path(OUTPUT_CONFIG["output_dir"])

SONNET_MODEL = "claude-sonnet-4-5"

# Max chars we pull from the tailored CV to give Claude context
MAX_CV_CHARS = 3000


class CoverLetterWriter:
    """
    Generates a job-specific, human-sounding cover letter as a .docx file.

    Usage:
        writer = CoverLetterWriter()
        path = writer.write(job, company_profile, cv_path)
    """

    def __init__(self):
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise EnvironmentError("ANTHROPIC_API_KEY not set in .env")
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model  = SONNET_MODEL
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # write() — public entry point
    # ------------------------------------------------------------------

    def write(self, job: dict, company_profile: Optional[dict] = None,
              cv_path: Optional[Path] = None) -> Path:
        """
        Generate a tailored cover letter .docx for the given job and company.

        Args:
            job             : dict from job_finder — needs title, company,
                              description, location, url
            company_profile : dict from company_researcher — used to personalise
                              paragraph 1 with real company-specific facts
            cv_path         : Path to the tailored CV .docx produced by
                              cv_customizer — we read its text to mirror
                              the same keywords Claude chose there

        Returns:
            Path to the generated cover letter .docx
        """
        company_name = job.get("company") or (
            company_profile.get("name") if company_profile else "the company"
        )
        job_title = job.get("title", "Software Engineer")

        print(f"\n[cover_letter] Writing cover letter for: {job_title} @ {company_name}")

        # Step 1 — load template (used for fallback and contact hydration)
        template = self._load_base_template()

        # Step 2 — read tailored CV text for keyword context
        cv_text = self._read_cv_text(cv_path) if cv_path else ""

        # Step 3 — call Claude
        letter_body = self._call_claude(
            job=job,
            company_profile=company_profile or {},
            cv_text=cv_text,
            template=template,
        )

        # Step 4 — build the full letter (salutation + body + sign-off)
        full_letter = self._assemble_letter(letter_body, job, company_profile or {})

        # Step 5 — determine output filename
        filename = self._make_filename(company_name, job_title)

        # Step 6 — render .docx
        output_path = self._write_docx(full_letter, filename)

        print(f"[cover_letter] Saved: {output_path}")
        return output_path

    # ------------------------------------------------------------------
    # _load_base_template
    # ------------------------------------------------------------------

    def _load_base_template(self) -> str:
        """
        Load the plain-text cover letter template from templates/.

        The template acts as:
          (a) structural guidance for the Claude prompt
          (b) a fallback document if Claude fails

        Returns the raw template string with {{PLACEHOLDERS}} intact.
        """
        try:
            return BASE_TEMPLATE_PATH.read_text(encoding="utf-8")
        except Exception as e:
            print(f"[cover_letter] Could not load template: {e}")
            return ""

    # ------------------------------------------------------------------
    # _read_cv_text
    # ------------------------------------------------------------------

    def _read_cv_text(self, cv_path: Path) -> str:
        """
        Extract plain text from the tailored CV .docx that cv_customizer produced.

        Why we read the CV:
          Claude already chose specific keywords and action verbs when writing
          the CV. The cover letter should echo the same language so both
          documents feel cohesive — an ATS and a recruiter reading both will
          see consistent terminology.

        We cap at MAX_CV_CHARS to keep the overall Claude prompt under the
        token limit while still giving enough context.

        Returns empty string if the file is missing or unreadable.
        """
        if not cv_path or not Path(cv_path).exists():
            return ""
        try:
            doc   = Document(str(cv_path))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            text  = "\n".join(lines)
            return text[:MAX_CV_CHARS]
        except Exception as e:
            print(f"[cover_letter] Could not read CV file ({cv_path}): {e}")
            return ""

    # ------------------------------------------------------------------
    # _build_prompt
    # ------------------------------------------------------------------

    def _build_prompt(self, job: dict, company_profile: dict,
                      cv_text: str, template: str) -> str:
        """
        Construct the Claude prompt for cover letter writing.

        Prompt structure:
          [Role description] → [Candidate facts] → [Job details] →
          [Company research] → [CV keywords] → [Template] → [Rules]

        Key design choices:
          - We pass the base template so Claude understands the expected
            structure and can see the placeholder labels.
          - We give hard word-count and paragraph-count constraints.
          - We tell Claude exactly what each paragraph must achieve.
          - We ask for ONLY the 3 body paragraphs — no salutation, no
            sign-off. We assemble those separately in _assemble_letter()
            so we control the exact formatting.

        The "Professional but human — not robotic" instruction is important:
          without it, Claude defaults to stiff, formal cover-letter clichés
          like "I am writing to express my interest in...".
        """
        cp = company_profile

        company_block = f"""
Company name:    {cp.get('name', job.get('company', 'Unknown'))}
Company size:    {cp.get('company_size', 'Unknown')}
Funding stage:   {cp.get('funding_stage', 'Unknown')}
Tech stack:      {', '.join(cp.get('tech_stack', [])) or 'Unknown'}
Overview:        {cp.get('overview', '')}
Culture notes:   {cp.get('culture_notes', '')}
Why apply (from research): {cp.get('why_apply', '')}
Fit score:       {cp.get('fit_score', 'N/A')}/10
Red flags:       {', '.join(cp.get('red_flags', [])) or 'None'}
""".strip()

        return f"""
You are writing a cover letter for Tanzil Ahmed, a Full-Stack Java + MERN developer
with Data Engineering experience (Kafka, PySpark, GCP, Azure), based in Bengaluru.

CANDIDATE CONTACT INFO (do NOT include in output — handled separately):
  Email:    {os.environ.get('CANDIDATE_EMAIL', '')}
  Phone:    {os.environ.get('CANDIDATE_PHONE', '')}
  LinkedIn: {os.environ.get('CANDIDATE_LINKEDIN', '')}
  GitHub:   {os.environ.get('CANDIDATE_GITHUB', '')}

============================
JOB DETAILS:
============================
Title:    {job.get('title', '')}
Company:  {job.get('company', '')}
Location: {job.get('location', '')}

Job Description:
{job.get('description', '(not available)')}

============================
COMPANY RESEARCH:
============================
{company_block}

============================
TAILORED CV CONTENT (keywords to echo):
============================
{cv_text or "(CV not available — use job description keywords instead)"}

============================
BASE TEMPLATE STRUCTURE (for reference):
============================
{template}

============================
YOUR TASK:
============================
Write EXACTLY 3 paragraphs. Output ONLY the 3 paragraphs — no date, no
address block, no salutation, no sign-off, no subject line, no labels.
Just the 3 paragraphs separated by a blank line.

PARAGRAPH 1 — Why THIS company (60-80 words):
  - Open with a specific, genuine hook about this company — NOT a generic
    "I am writing to express my interest" opener.
  - Reference one concrete fact from the company research (tech stack,
    mission, product, culture, recent news) that genuinely excites Tanzil.
  - End by stating the role being applied for and Tanzil's core identity
    as a developer in one natural sentence.
  - Tone: enthusiastic but grounded, like a real person wrote it.

PARAGRAPH 2 — Why Tanzil is the perfect fit (140-160 words):
  - Mirror 3-4 exact keywords or phrases from the job description.
  - Connect each keyword to a specific skill or achievement from the CV.
  - Include at least one quantified result if the CV provides one.
  - Do NOT list skills robotically ("I have X, I have Y") — weave them
    into sentences that tell a story of what Tanzil built and achieved.

PARAGRAPH 3 — Confident close with call to action (50-70 words):
  - Express forward-looking enthusiasm — what Tanzil hopes to contribute,
    not just what he wants to gain.
  - Include a clear, direct call to action (interview request).
  - Mention notice period: Immediate / 15 days.
  - End on a confident, warm note — not desperate or over-eager.

TOTAL: 350 words maximum.
TONE: Professional but human. Confident, not arrogant. Specific, not generic.
DO NOT use these clichés: "I am writing to", "passion for", "team player",
"fast learner", "hard worker", "results-driven", "I would be a great fit".
"""

    # ------------------------------------------------------------------
    # _call_claude
    # ------------------------------------------------------------------

    def _call_claude(self, job: dict, company_profile: dict,
                     cv_text: str, template: str) -> str:
        """
        Send the prompt to Claude Sonnet and return only the 3 body paragraphs.

        Why Sonnet (not Haiku):
          Cover letters require genuine creativity and nuanced language.
          Haiku is fast but produces noticeably stiffer, more generic prose.
          Since we produce one cover letter per application, the cost
          difference is trivial and the quality difference is meaningful.

        Fallback:
          If Claude fails, we call _build_fallback() which fills the base
          template with known facts — a mediocre-but-complete letter is
          better than an empty file that stalls the pipeline.
        """
        prompt = self._build_prompt(job, company_profile, cv_text, template)
        try:
            message = self.client.messages.create(
                model=self.model,
                max_tokens=1024,
                messages=[{"role": "user", "content": prompt}],
            )
            body = message.content[0].text.strip()
            print(f"[cover_letter] Claude returned {len(body)} chars ({len(body.split())} words)")
            return body

        except Exception as e:
            print(f"[cover_letter] Claude call failed ({e}) — using fallback template")
            return self._build_fallback(job, company_profile)

    # ------------------------------------------------------------------
    # _build_fallback
    # ------------------------------------------------------------------

    def _build_fallback(self, job: dict, company_profile: dict) -> str:
        """
        Build a minimal cover letter body when Claude is unavailable.

        Substitutes known values (company name, job title, tech stack) into
        fixed sentences so the result is at least specific to this application.
        Not impressive, but complete and truthful.
        """
        company    = job.get("company", "your company")
        title      = job.get("title", "this role")
        tech_stack = ", ".join((company_profile.get("tech_stack") or [])[:4]) or "your tech stack"

        return (
            f"I was excited to come across the {title} opening at {company}. "
            f"Your work in {company_profile.get('overview', 'this space')[:120]} "
            f"aligns closely with where I want to take my career as a Full-Stack Java "
            f"and MERN developer with Data Engineering experience.\n\n"
            f"My background covers the core technologies your team relies on, including "
            f"{tech_stack}. I have hands-on experience building scalable backend systems "
            f"with Java and Spring Boot, designing real-time data pipelines with Kafka and "
            f"PySpark, and delivering full-stack features with React on GCP and Azure. I am "
            f"comfortable in agile environments and have contributed to CI/CD pipelines, "
            f"microservices architectures, and cross-functional teams throughout my work.\n\n"
            f"I would welcome the opportunity to discuss how my skills align with your team's "
            f"goals. I am available for an interview at your convenience and can join within "
            f"15 days. Thank you for considering my application."
        )

    # ------------------------------------------------------------------
    # _assemble_letter
    # ------------------------------------------------------------------

    def _assemble_letter(self, body: str, job: dict,
                         company_profile: dict) -> dict:
        """
        Wrap Claude's 3 body paragraphs in the full letter structure.

        Returns a dict of named blocks that _write_docx() renders separately,
        giving us fine-grained control over the spacing and font of each part:

          {
            "date":        "31 March 2026",
            "subject":     "Application for Associate Software Engineer — Tanzil Ahmed",
            "salutation":  "Dear Hiring Manager,",
            "paragraphs":  ["Para 1 text", "Para 2 text", "Para 3 text"],
            "signoff":     "Warm regards,",
            "name":        "Tanzil Ahmed",
            "contact":     "+91-8753909446  |  tanzilahmed37@gmail.com  |  ...",
          }

        Why split into a dict rather than one big string:
          Each part gets different formatting in the docx (bold subject,
          indented paragraphs, right-aligned sign-off, etc.) — we need
          them separate to apply that formatting cleanly.
        """
        company_name = job.get("company") or company_profile.get("name", "the company")
        job_title    = job.get("title", "Software Engineer")

        # Split Claude's output into individual paragraphs (separated by blank lines)
        raw_paragraphs = [p.strip() for p in re.split(r"\n{2,}", body) if p.strip()]

        # Ensure we always have exactly 3 paragraphs (pad with empty if Claude
        # returned fewer, which can happen with very short responses)
        while len(raw_paragraphs) < 3:
            raw_paragraphs.append("")

        # Build contact footer line
        contact_parts = []
        phone    = os.environ.get("CANDIDATE_PHONE", "")
        email    = os.environ.get("CANDIDATE_EMAIL", "")
        linkedin = os.environ.get("CANDIDATE_LINKEDIN", "")
        github   = os.environ.get("CANDIDATE_GITHUB", "")
        for part in [phone, email, linkedin, github]:
            if part:
                contact_parts.append(part)

        return {
            "date":       date.today().strftime("%-d %B %Y") if os.name != "nt"
                          else date.today().strftime("%d %B %Y").lstrip("0"),
            "subject":    f"Application for {job_title} — Tanzil Ahmed",
            "salutation": "Dear Hiring Manager,",
            "paragraphs": raw_paragraphs[:3],
            "signoff":    "Warm regards,",
            "name":       "Tanzil Ahmed",
            "contact":    "  |  ".join(contact_parts),
            # Metadata — rendered as a grey header block at the top
            "apply_url":  job.get("url", ""),
            "company":    company_name,
            "job_title":  job_title,
            "fit_score":  company_profile.get("fit_score", "?"),
            "date_found": date.today().strftime("%d %B %Y"),
        }

    # ------------------------------------------------------------------
    # _make_filename
    # ------------------------------------------------------------------

    def _make_filename(self, company_name: str, job_title: str) -> str:
        """
        Build a safe, sortable filename for the cover letter .docx.

        Format: cover_[company]_[role]_[YYYY-MM-DD].docx

        Uses the same slugify logic as cv_customizer so both files for the
        same application sort together in the output/ folder:
          cv_thoughtworks_associate_software_engineer_2026-03-31.docx
          cover_thoughtworks_associate_software_engineer_2026-03-31.docx
        """
        def slugify(s: str) -> str:
            s = s.lower().strip()
            s = re.sub(r"[^\w\s-]", "", s)
            s = re.sub(r"[\s-]+", "_", s)
            return s[:30]

        company_slug = slugify(company_name)
        role_slug    = slugify(job_title)
        today        = date.today().isoformat()
        return f"cover_{company_slug}_{role_slug}_{today}.docx"

    # ------------------------------------------------------------------
    # _write_docx
    # ------------------------------------------------------------------

    def _write_docx(self, letter: dict, filename: str) -> Path:
        """
        Render the assembled cover letter into a clean, professional .docx.

        Layout decisions:
          - Same margins and font family as the CV (Calibri) for brand consistency
          - Date: right-aligned, 10pt
          - Subject line: bold, 11pt — makes it easy to scan in an inbox
          - Salutation: normal, with extra space below
          - Body paragraphs: 11pt, justified, 6pt spacing between them —
            the slight paragraph spacing reads better than double-spacing
          - Sign-off and name: left-aligned
          - Contact line: 9pt, grey — subtle but present

        Why NOT use a table for the header layout:
          Some ATS systems strip table content. Plain paragraphs are safer
          for automated parsing while still looking professional.
        """
        doc = Document()

        # ── Page margins ────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Pt(54)   # 0.75"
            section.bottom_margin = Pt(54)
            section.left_margin   = Pt(72)   # 1"
            section.right_margin  = Pt(72)

        # ── Base style ──────────────────────────────────────────────────
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after  = Pt(6)

        # ── Job metadata block (apply link, fit score, date) ───────────
        meta_lines = [
            f"APPLY HERE:  {letter.get('apply_url', '(no URL)')}",
            f"Company:     {letter.get('company', '?')}",
            f"Role:        {letter.get('job_title', '?')}",
            f"Fit Score:   {letter.get('fit_score', '?')}/10",
            f"Date Found:  {letter.get('date_found', '')}",
        ]
        for line in meta_lines:
            p   = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size      = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            run.font.name      = "Calibri"
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
        div = doc.add_paragraph()
        div.add_run("─" * 68).font.size = Pt(7)
        div.paragraph_format.space_before = Pt(2)
        div.paragraph_format.space_after  = Pt(10)

        # ── Date — right-aligned ────────────────────────────────────────
        self._add_para(doc, letter["date"], align=WD_ALIGN_PARAGRAPH.RIGHT,
                       size=10, space_after=12)

        # ── Blank spacer ────────────────────────────────────────────────
        self._add_para(doc, "", space_after=4)

        # ── Subject line — bold ─────────────────────────────────────────
        self._add_para(doc, letter["subject"], bold=True, size=11, space_after=12)

        # ── Salutation ──────────────────────────────────────────────────
        self._add_para(doc, letter["salutation"], space_after=10)

        # ── Body paragraphs ─────────────────────────────────────────────
        for para_text in letter["paragraphs"]:
            if para_text:
                self._add_para(doc, para_text,
                               align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               size=11, space_after=10)

        # ── Sign-off ────────────────────────────────────────────────────
        self._add_para(doc, "", space_after=4)
        self._add_para(doc, letter["signoff"], space_after=2)

        # ── Name — bold ─────────────────────────────────────────────────
        self._add_para(doc, letter["name"], bold=True, size=11, space_after=4)

        # ── Contact line — small and grey ───────────────────────────────
        if letter["contact"]:
            p = doc.add_paragraph()
            run = p.add_run(letter["contact"])
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)   # mid-grey
            p.paragraph_format.space_after = Pt(0)

        # ── Save ────────────────────────────────────────────────────────
        output_path = OUTPUT_DIR / filename
        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # _add_para — internal docx helper
    # ------------------------------------------------------------------

    def _add_para(self, doc: Document, text: str,
                  bold: bool = False,
                  size: float = 11,
                  align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
                  space_after: int = 6) -> None:
        """
        Add a single paragraph to the document with consistent styling.

        Centralising paragraph creation here means every call site specifies
        only what differs from the default — font size, bold, alignment,
        spacing — without repeating boilerplate docx API calls.

        Args:
            doc        : the Document being built
            text       : paragraph content (empty string = blank spacer line)
            bold       : whether the run is bold
            size       : font size in points
            align      : paragraph alignment constant
            space_after: space after the paragraph in points
        """
        p   = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)

        if text:
            run = p.add_run(text)
            run.bold      = bold
            run.font.size = Pt(size)
            run.font.name = "Calibri"
