"""
tools/cv_customizer.py — Tailors Tanzil's base CV to a specific job + company.

Pipeline:
  1. Load base CV from templates/cv_base.txt
  2. Hydrate contact placeholders from .env
  3. Build a detailed prompt: base CV + job description + company profile
  4. Call Claude Sonnet (best model — CV quality is the most important output)
  5. Parse Claude's response into structured sections
  6. Render a properly formatted .docx with python-docx
  7. Return the file path so the tracker and orchestrator can record it

Fallback: if Claude fails at any point, the base CV is written as-is to .docx
so the pipeline never stalls.
"""

import os
import re
from datetime import date
from pathlib import Path
from typing import Optional

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv

from config import CANDIDATE_PROFILE, OUTPUT_CONFIG

load_dotenv()

# ---------------------------------------------------------------------------
# Paths and model
# ---------------------------------------------------------------------------
BASE_CV_PATH = Path("templates/cv_base.txt")
OUTPUT_DIR   = Path(OUTPUT_CONFIG["output_dir"])

# Sonnet for CV — highest reasoning quality needed here
SONNET_MODEL = "claude-sonnet-4-5"

# Section heading markers Claude must use in its response
SECTION_MARKERS = [
    "PROFESSIONAL SUMMARY",
    "SKILLS",
    "EXPERIENCE",
    "PROJECTS",
    "EDUCATION",
    "CERTIFICATIONS",
    "ADDITIONAL",
]


class CVCustomizer:
    """
    Generates a job-specific, ATS-optimised CV as a .docx file.

    Usage:
        customizer = CVCustomizer()
        path = customizer.customise(job, company_profile)
    """

    def __init__(self):
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise EnvironmentError("ANTHROPIC_API_KEY not set in .env")
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model  = SONNET_MODEL
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # customise() — public entry point
    # ------------------------------------------------------------------

    def customise(self, job: dict, company_profile: Optional[dict] = None) -> Path:
        """
        Produce a tailored CV .docx for the given job and company.

        Args:
            job             : dict from job_finder — needs at minimum:
                                title, company, description, url, location
            company_profile : dict from company_researcher (optional).
                              If provided, the summary and skills sections
                              are tuned to the company's actual tech stack.

        Returns:
            Path to the generated .docx file.
        """
        company_name = job.get("company") or (
            company_profile.get("name") if company_profile else "Company"
        )
        job_title = job.get("title", "Software Engineer")

        print(f"\n[cv_customizer] Tailoring CV for: {job_title} @ {company_name}")

        # Step 1 — load and hydrate the base CV
        base_cv = self._load_base_cv()

        # Step 2 — ask Claude to tailor it
        tailored_text = self._call_claude(
            base_cv=base_cv,
            job=job,
            company_profile=company_profile or {},
        )

        # Step 3 — determine output filename
        filename = self._make_filename(company_name, job_title)

        # Step 4 — render .docx
        output_path = self._write_docx(tailored_text, filename, job_title,
                                       company_name, job, company_profile or {})

        print(f"[cv_customizer] Saved: {output_path}")
        return output_path

    # ------------------------------------------------------------------
    # _load_base_cv
    # ------------------------------------------------------------------

    def _load_base_cv(self) -> str:
        """
        Load the plain-text base CV and substitute contact placeholders
        with real values from .env / config.

        Why substitute here rather than in the template:
          The template stores placeholders so it stays shareable and
          version-controllable without exposing personal data. We inject
          the real values at runtime from .env.
        """
        text = BASE_CV_PATH.read_text(encoding="utf-8")

        replacements = {
            "[your_email]":   os.environ.get("CANDIDATE_EMAIL", CANDIDATE_PROFILE.get("email", "")),
            "[your_phone]":   os.environ.get("CANDIDATE_PHONE", CANDIDATE_PROFILE.get("phone", "")),
            "[linkedin_url]": os.environ.get("CANDIDATE_LINKEDIN", CANDIDATE_PROFILE.get("linkedin", "")),
            "[github_url]":   os.environ.get("CANDIDATE_GITHUB", CANDIDATE_PROFILE.get("github", "")),
        }
        for placeholder, value in replacements.items():
            text = text.replace(placeholder, value)

        return text

    # ------------------------------------------------------------------
    # _build_prompt
    # ------------------------------------------------------------------

    def _build_prompt(self, base_cv: str, job: dict,
                      company_profile: dict) -> str:
        """
        Construct the Claude prompt.

        Design decisions:
          - We give Claude the full base CV, full job description, and the
            company profile so it has maximum context.
          - We specify exact section headings it must use — this makes
            _parse_sections() reliable.
          - We instruct it to maximise ATS keyword overlap without fabricating.
          - We ask for plain text output (no markdown) because python-docx
            does not render markdown — we handle all formatting ourselves.

        The prompt follows the structure:
          [Role] → [Input data] → [Rules] → [Output format]
        """
        # Flatten company profile into a readable block
        cp = company_profile
        company_block = f"""
Company name:    {cp.get('name', job.get('company', 'Unknown'))}
Company size:    {cp.get('company_size', 'Unknown')}
Tech stack:      {', '.join(cp.get('tech_stack', [])) or 'Unknown'}
Culture notes:   {cp.get('culture_notes', '')}
Funding stage:   {cp.get('funding_stage', 'Unknown')}
Why apply:       {cp.get('why_apply', '')}
Fit score:       {cp.get('fit_score', 'N/A')}/10
""".strip()

        return f"""
You are an expert CV writer and ATS optimisation specialist.

Your task is to tailor Tanzil Ahmed's CV for a specific job application.
Rewrite the CV so it:
  1. Mirrors the exact keywords and phrases from the job description
     (ATS systems do literal keyword matching — this is critical)
  2. Reorders skills to put the most relevant ones first
  3. Rewrites the professional summary to specifically address this company
     and role (use the company's language and values)
  4. Strengthens bullet points in Experience and Projects with the job's
     preferred action verbs and technical terms
  5. Stays 100% truthful — only REFRAME existing facts, never fabricate
  6. Removes or de-emphasises skills irrelevant to this role

============================
CANDIDATE'S BASE CV:
============================
{base_cv}

============================
JOB DETAILS:
============================
Title:       {job.get('title', '')}
Company:     {job.get('company', '')}
Location:    {job.get('location', '')}
URL:         {job.get('url', '')}

Job Description:
{job.get('description', '(description not available)')}

============================
COMPANY PROFILE (from research):
============================
{company_block}

============================
OUTPUT RULES — READ CAREFULLY:
============================
- Output ONLY the tailored CV text. No introduction, no explanation, no markdown.
- Use EXACTLY these section headings on their own line, in ALL CAPS:
    TANZIL AHMED          ← name first (not a section heading)
    PROFESSIONAL SUMMARY
    SKILLS
    EXPERIENCE
    PROJECTS
    EDUCATION
    CERTIFICATIONS
    ADDITIONAL
- Keep the contact line (email | phone | linkedin | github) directly under the name.
- For skills: group into labelled rows like "Backend : Java, Spring Boot, ..."
  Put the most relevant skill groups for THIS job first.
- For bullet points: start every bullet with a strong action verb.
  Each bullet must include at least one keyword from the job description.
- Keep the CV to a maximum of 2 pages of content.
- Do NOT add any section the base CV does not have.
- Do NOT add horizontal rules or decorative characters — plain text only.
"""

    # ------------------------------------------------------------------
    # _call_claude
    # ------------------------------------------------------------------

    def _call_claude(self, base_cv: str, job: dict,
                     company_profile: dict) -> str:
        """
        Send the prompt to Claude Sonnet and return the tailored CV text.

        Why Sonnet here:
          CV tailoring requires nuanced reasoning — understanding which of
          Tanzil's experiences best map to each requirement, choosing the
          right action verbs, writing a compelling summary. Haiku is too
          shallow for this. Sonnet gives a measurably better result.

        Fallback:
          If the API call fails for any reason (rate limit, network, etc.)
          we return the base CV unchanged. The docx will still be written —
          it just won't be customised. The caller is notified via print.
        """
        prompt = self._build_prompt(base_cv, job, company_profile)
        try:
            message = self.client.messages.create(
                model=self.model,
                max_tokens=4096,   # CVs can be long — give Claude room
                messages=[{"role": "user", "content": prompt}],
            )
            tailored = message.content[0].text.strip()
            print(f"[cv_customizer] Claude returned {len(tailored)} chars")
            return tailored

        except Exception as e:
            print(f"[cv_customizer] Claude call failed ({e}) — using base CV as fallback")
            return base_cv

    # ------------------------------------------------------------------
    # _parse_sections
    # ------------------------------------------------------------------

    def _parse_sections(self, cv_text: str) -> dict:
        """
        Split the flat CV text Claude produced into a dict of sections.

        Why parse rather than dump the whole text into one paragraph:
          python-docx gives us full control over font sizes, bold headings,
          bullet spacing, etc. To apply that formatting we need to know
          which text belongs to which section.

        Strategy:
          We scan line-by-line. When we see a line that exactly matches
          one of our SECTION_MARKERS (case-insensitive), we start a new
          section bucket. Everything else gets appended to the current bucket.

        Returns a dict like:
          {
            "header":               "TANZIL AHMED\nBengaluru...",
            "PROFESSIONAL SUMMARY": "Full-Stack Java...",
            "SKILLS":               "Backend : Java...",
            ...
          }
        """
        sections: dict = {"header": []}
        current = "header"

        for line in cv_text.splitlines():
            stripped = line.strip()
            upper    = stripped.upper()

            if upper in SECTION_MARKERS:
                current = upper
                sections[current] = []
            else:
                sections.setdefault(current, [])
                sections[current].append(line)

        # Join each bucket into a single string, strip leading/trailing blanks
        return {k: "\n".join(v).strip() for k, v in sections.items()}

    # ------------------------------------------------------------------
    # _make_filename
    # ------------------------------------------------------------------

    def _make_filename(self, company_name: str, job_title: str) -> str:
        """
        Build a safe, sortable filename for the output .docx.

        Format: cv_[company]_[role]_[YYYY-MM-DD].docx

        We slugify both strings: lowercase, spaces → underscores,
        drop everything that isn't alphanumeric or an underscore.
        This keeps filenames safe on all OSes and easy to sort by date.

        Example: cv_thoughtworks_associate_software_engineer_2026-03-31.docx
        """
        def slugify(s: str) -> str:
            s = s.lower().strip()
            s = re.sub(r"[^\w\s-]", "", s)
            s = re.sub(r"[\s-]+", "_", s)
            return s[:30]   # cap length

        company_slug = slugify(company_name)
        role_slug    = slugify(job_title)
        today        = date.today().isoformat()   # e.g. "2026-03-31"
        return f"cv_{company_slug}_{role_slug}_{today}.docx"

    # ------------------------------------------------------------------
    # _write_docx
    # ------------------------------------------------------------------

    def _write_docx(self, cv_text: str, filename: str,
                    job_title: str, company_name: str,
                    job: dict = None, company_profile: dict = None) -> Path:
        """
        Render the tailored CV text into a properly formatted .docx file.

        Formatting decisions:
          - Name: 16pt bold, centred
          - Contact line: 10pt, centred
          - Section headings: 12pt bold, left-aligned, with a bottom border
            simulated by a full-width underline paragraph
          - Body text: 11pt, left-aligned
          - Bullet lines (starting with •): 11pt, indented 0.3"
          - Skill rows ("Backend : Java..."): 10.5pt, label in bold
          - Blank lines between sections preserved for readability

        Why python-docx instead of just saving a .txt:
          Recruiters open Word documents. A .docx lets us control fonts,
          spacing, and layout — things that matter when the CV is printed
          or pasted into an ATS upload field.
        """
        doc = Document()

        # ── Global page margins: narrow to fit more content ────────────
        for section in doc.sections:
            section.top_margin    = Pt(36)   # 0.5"
            section.bottom_margin = Pt(36)
            section.left_margin   = Pt(54)   # 0.75"
            section.right_margin  = Pt(54)

        # Remove default paragraph spacing
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after  = Pt(2)

        # ── Job metadata block (apply link, fit score, date) ───────────
        job             = job or {}
        company_profile = company_profile or {}
        self._add_metadata_block(doc, job, company_profile)

        # ── Parse the CV text into sections ────────────────────────────
        sections = self._parse_sections(cv_text)

        # ── Header block (name + contact) ──────────────────────────────
        self._add_header(doc, sections.get("header", ""))

        # ── Body sections in order ─────────────────────────────────────
        section_order = [
            "PROFESSIONAL SUMMARY",
            "SKILLS",
            "EXPERIENCE",
            "PROJECTS",
            "EDUCATION",
            "CERTIFICATIONS",
            "ADDITIONAL",
        ]
        for heading in section_order:
            content = sections.get(heading, "")
            if not content:
                continue
            self._add_section_heading(doc, heading.title())
            self._add_section_body(doc, content)

        # ── Save ────────────────────────────────────────────────────────
        output_path = OUTPUT_DIR / filename
        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # docx helper methods
    # ------------------------------------------------------------------

    def _add_metadata_block(self, doc: Document, job: dict,
                            company_profile: dict) -> None:
        """
        Write a small grey metadata block at the very top of the document.

        This block tells Tanzil exactly where to apply and what the agent
        scored the role — without affecting the CV content below it.
        The grey colour keeps it visually separate from the actual CV.
        """
        lines = [
            f"APPLY HERE:  {job.get('url', '(no URL)')}",
            f"Company:     {job.get('company') or company_profile.get('name', '?')}",
            f"Role:        {job.get('title', '?')}",
            f"Fit Score:   {company_profile.get('fit_score', '?')}/10",
            f"Date Found:  {date.today().strftime('%d %B %Y')}",
        ]
        for i, line in enumerate(lines):
            p   = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size      = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)   # mid-grey
            run.font.name      = "Calibri"
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)

        # Thin divider after the block
        div = doc.add_paragraph()
        div.add_run("─" * 68).font.size = Pt(7)
        div.paragraph_format.space_before = Pt(2)
        div.paragraph_format.space_after  = Pt(8)

    def _add_header(self, doc: Document, header_text: str) -> None:
        """
        Write the name and contact block at the top of the document.

        The first non-empty line is treated as the candidate name (large + bold).
        All subsequent lines are the contact info (small + centred).
        """
        lines = [l for l in header_text.splitlines()]
        if not lines:
            return

        # Name — first line
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(lines[0].strip())
        run.bold      = True
        run.font.size = Pt(16)
        name_para.paragraph_format.space_after = Pt(2)

        # Contact lines — everything after the name
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(1)

        # Visual separator after header
        doc.add_paragraph()

    def _add_section_heading(self, doc: Document, title: str) -> None:
        """
        Add a bold, slightly larger section heading with a thin divider line below.

        The divider is a separate paragraph of underscores at 7pt — visually
        similar to the ━━━ in the text template but compatible with all .docx viewers.
        """
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold         = True
        run.font.size    = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # dark blue — professional look
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(0)

        # Thin divider
        div = doc.add_paragraph()
        div.add_run("─" * 68).font.size = Pt(7)
        div.paragraph_format.space_before = Pt(0)
        div.paragraph_format.space_after  = Pt(4)

    def _add_section_body(self, doc: Document, content: str) -> None:
        """
        Write the body text of a section, applying special formatting per line:

          • lines starting with "•" or "-"  → bullet indented paragraph
          • lines matching "Label : value"   → label bolded (skills rows)
          • blank lines                      → small spacer paragraph
          • everything else                  → normal body paragraph

        This gives the CV proper visual hierarchy without needing Heading styles.
        """
        for line in content.splitlines():
            stripped = line.strip()

            # ── Empty line → thin spacer ────────────────────────────
            if not stripped:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                continue

            # ── Bullet point ─────────────────────────────────────────
            if stripped.startswith(("•", "-", "*")):
                bullet_text = stripped.lstrip("•-* ").strip()
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(bullet_text)
                run.font.size = Pt(10.5)
                p.paragraph_format.left_indent   = Pt(18)
                p.paragraph_format.space_after   = Pt(2)
                p.paragraph_format.space_before  = Pt(0)
                continue

            # ── Skill row: "Label : value1, value2" ──────────────────
            # Match lines like "Backend : Java, Spring Boot, ..."
            skill_match = re.match(r"^([A-Za-z /]+)\s*:\s*(.+)$", stripped)
            if skill_match:
                p = doc.add_paragraph()
                label_run = p.add_run(skill_match.group(1).strip() + " : ")
                label_run.bold      = True
                label_run.font.size = Pt(10.5)
                value_run = p.add_run(skill_match.group(2).strip())
                value_run.font.size = Pt(10.5)
                p.paragraph_format.space_after = Pt(2)
                continue

            # ── Regular body line ─────────────────────────────────────
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(2)
